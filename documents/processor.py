"""
Document Processing System
Extract text and data from PDF, DOCX, XLSX files
"""

from typing import Dict, List, Optional, BinaryIO
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl
from datetime import datetime
from uuid import UUID
from sqlalchemy.ext.asyncio import AsyncSession
import json
import re

from models import ProcessedDocument
from loguru import logger


class DocumentProcessor:
    """
    Process various document types
    Extract text, tables, metadata
    """

    def __init__(self):
        self.supported_types = ['pdf', 'docx', 'xlsx', 'txt']

    async def process_document(
        self,
        db: AsyncSession,
        file: BinaryIO,
        filename: str,
        user_id: UUID,
        document_type: Optional[str] = None,
        metadata: Optional[Dict] = None
    ) -> Dict:
        """
        Process document and extract content

        Args:
            db: Database session
            file: File object
            filename: Original filename
            user_id: User uploading document
            document_type: Type hint (estimates, policies, reports, etc.)
            metadata: Additional metadata

        Returns:
            Dict with extracted content and analysis
        """
        try:
            # Determine file type
            file_extension = Path(filename).suffix.lower().replace('.', '')

            if file_extension not in self.supported_types:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Process based on type
            if file_extension == 'pdf':
                result = await self._process_pdf(file, filename)
            elif file_extension == 'docx':
                result = await self._process_docx(file, filename)
            elif file_extension == 'xlsx':
                result = await self._process_xlsx(file, filename)
            elif file_extension == 'txt':
                result = await self._process_txt(file, filename)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Extract key information
            extracted_info = self._extract_key_information(result['text'])

            # Save to database
            processed_doc = ProcessedDocument(
                user_id=user_id,
                filename=filename,
                file_type=file_extension,
                document_type=document_type or "unknown",
                extracted_text=result['text'],
                extracted_data=result.get('structured_data', {}),
                metadata={
                    **(metadata or {}),
                    "page_count": result.get('page_count'),
                    "word_count": len(result['text'].split()),
                    "extracted_info": extracted_info,
                    "processing_date": datetime.utcnow().isoformat()
                }
            )

            db.add(processed_doc)
            await db.flush()

            logger.info(f"Processed document: {filename} ({file_extension}) for user {user_id}")

            return {
                "document_id": str(processed_doc.id),
                "filename": filename,
                "file_type": file_extension,
                "text": result['text'],
                "structured_data": result.get('structured_data', {}),
                "extracted_info": extracted_info,
                "page_count": result.get('page_count'),
                "word_count": len(result['text'].split())
            }

        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}", exc_info=True)
            raise

    async def _process_pdf(self, file: BinaryIO, filename: str) -> Dict:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text_parts = []

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                text_parts.append(f"--- Page {page_num} ---\n{text}\n")

            full_text = "\n".join(text_parts)

            logger.info(f"Extracted {len(pdf_reader.pages)} pages from PDF: {filename}")

            return {
                "text": full_text,
                "page_count": len(pdf_reader.pages),
                "metadata": pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
            }

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

    async def _process_docx(self, file: BinaryIO, filename: str) -> Dict:
        """Extract text from DOCX file"""
        try:
            doc = Document(file)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)

            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)

            logger.info(f"Extracted {len(paragraphs)} paragraphs and {len(tables_data)} tables from DOCX: {filename}")

            return {
                "text": full_text,
                "page_count": len(doc.sections),
                "structured_data": {
                    "tables": tables_data,
                    "paragraph_count": len(paragraphs)
                }
            }

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

    async def _process_xlsx(self, file: BinaryIO, filename: str) -> Dict:
        """Extract data from XLSX file"""
        try:
            workbook = openpyxl.load_workbook(file, data_only=True)
            sheets_data = {}
            all_text = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                rows_data = []

                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    if any(cell is not None for cell in row):
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        rows_data.append(row_data)

                        # Build text representation
                        all_text.append(" | ".join(row_data))

                sheets_data[sheet_name] = rows_data

            full_text = "\n".join(all_text)

            logger.info(f"Extracted {len(sheets_data)} sheets from XLSX: {filename}")

            return {
                "text": full_text,
                "page_count": len(sheets_data),
                "structured_data": {
                    "sheets": sheets_data,
                    "sheet_names": list(sheets_data.keys())
                }
            }

        except Exception as e:
            logger.error(f"Error processing XLSX: {e}")
            raise

    async def _process_txt(self, file: BinaryIO, filename: str) -> Dict:
        """Extract text from TXT file"""
        try:
            content = file.read().decode('utf-8', errors='ignore')

            logger.info(f"Extracted text from TXT: {filename}")

            return {
                "text": content,
                "page_count": 1
            }

        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            raise

    def _extract_key_information(self, text: str) -> Dict:
        """
        Extract key information from document text
        Looks for common patterns in insurance documents

        Args:
            text: Document text

        Returns:
            Dict with extracted information
        """
        info = {
            "claim_numbers": [],
            "dates": [],
            "amounts": [],
            "addresses": [],
            "phone_numbers": [],
            "email_addresses": []
        }

        try:
            # Extract claim numbers (various formats)
            claim_patterns = [
                r'Claim\s*(?:Number|#|No\.?):\s*([A-Z0-9\-]+)',
                r'Claim\s*([A-Z0-9\-]{5,})',
                r'Policy\s*(?:Number|#):\s*([A-Z0-9\-]+)'
            ]
            for pattern in claim_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                info["claim_numbers"].extend(matches)

            # Extract dates (MM/DD/YYYY, MM-DD-YYYY)
            date_pattern = r'\b\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}\b'
            info["dates"] = re.findall(date_pattern, text)

            # Extract amounts (currency)
            amount_pattern = r'\$[\d,]+\.?\d{0,2}'
            info["amounts"] = re.findall(amount_pattern, text)

            # Extract email addresses
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            info["email_addresses"] = re.findall(email_pattern, text)

            # Extract phone numbers
            phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'
            info["phone_numbers"] = re.findall(phone_pattern, text)

            # Extract addresses (simplified - city, state zip pattern)
            address_pattern = r'[A-Z][a-z]+,\s+[A-Z]{2}\s+\d{5}'
            info["addresses"] = re.findall(address_pattern, text)

            # Deduplicate
            for key in info:
                info[key] = list(set(info[key]))[:10]  # Limit to 10 unique items

            logger.debug(f"Extracted key information: {len(info['claim_numbers'])} claims, {len(info['dates'])} dates, etc.")

        except Exception as e:
            logger.error(f"Error extracting key information: {e}")

        return info

    async def search_documents(
        self,
        db: AsyncSession,
        user_id: UUID,
        query: str,
        document_type: Optional[str] = None,
        limit: int = 20
    ) -> List[Dict]:
        """
        Search user's processed documents

        Args:
            db: Database session
            user_id: User ID
            query: Search query
            document_type: Filter by document type
            limit: Max results

        Returns:
            List of matching documents
        """
        try:
            from sqlalchemy import select, or_

            query_obj = select(ProcessedDocument).where(
                ProcessedDocument.user_id == user_id
            )

            if document_type:
                query_obj = query_obj.where(
                    ProcessedDocument.document_type == document_type
                )

            # Search in filename or extracted text
            search_term = f"%{query}%"
            query_obj = query_obj.where(
                or_(
                    ProcessedDocument.filename.ilike(search_term),
                    ProcessedDocument.extracted_text.ilike(search_term)
                )
            ).limit(limit)

            result = await db.execute(query_obj)
            documents = result.scalars().all()

            results = [
                {
                    "id": str(doc.id),
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "document_type": doc.document_type,
                    "processed_at": doc.processed_at.isoformat(),
                    "text_preview": doc.extracted_text[:500] if doc.extracted_text else None
                }
                for doc in documents
            ]

            logger.info(f"Found {len(results)} documents matching '{query}'")

            return results

        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return []

    async def analyze_estimate(self, text: str, structured_data: Optional[Dict] = None) -> Dict:
        """
        Analyze an estimate document
        Extract line items, totals, scope of work

        Args:
            text: Document text
            structured_data: Structured data if available (from XLSX)

        Returns:
            Dict with estimate analysis
        """
        try:
            analysis = {
                "line_items": [],
                "totals": {
                    "subtotal": None,
                    "tax": None,
                    "total": None
                },
                "scope_summary": "",
                "categories": []
            }

            # If structured data from Excel, parse that
            if structured_data and 'sheets' in structured_data:
                # Look for estimate sheet
                for sheet_name, rows in structured_data['sheets'].items():
                    if 'estimate' in sheet_name.lower() or 'line' in sheet_name.lower():
                        # Assuming first row is headers
                        if len(rows) > 1:
                            headers = rows[0]
                            for row in rows[1:]:
                                if len(row) >= 3:  # Assume at least description, qty, amount
                                    analysis["line_items"].append({
                                        "description": row[0] if len(row) > 0 else "",
                                        "quantity": row[1] if len(row) > 1 else "",
                                        "amount": row[2] if len(row) > 2 else ""
                                    })

            # Extract totals from text
            total_patterns = [
                (r'Total:\s*\$?([\d,]+\.?\d{0,2})', 'total'),
                (r'Subtotal:\s*\$?([\d,]+\.?\d{0,2})', 'subtotal'),
                (r'Tax:\s*\$?([\d,]+\.?\d{0,2})', 'tax'),
                (r'Grand\s+Total:\s*\$?([\d,]+\.?\d{0,2})', 'total')
            ]

            for pattern, field in total_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches and not analysis["totals"][field]:
                    analysis["totals"][field] = matches[0]

            logger.info(f"Analyzed estimate: {len(analysis['line_items'])} line items")

            return analysis

        except Exception as e:
            logger.error(f"Error analyzing estimate: {e}")
            return {}


# Global instance
document_processor = DocumentProcessor()
